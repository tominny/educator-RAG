import streamlit as st
import psycopg2
from langchain_openai import OpenAIEmbeddings
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain.chains import ConversationalRetrievalChain
from langchain_community.vectorstores import PGVector
from typing import List
import pandas as pd
import docx
import pdfplumber

CONN_STR = st.secrets["NEONDB_CONNECTION_STRING"]

def init_rag_db():
    """
    Ensure 'embeddings' table exists for storing document vectors (via pgvector).
    """
    with psycopg2.connect(CONN_STR) as conn:
        conn.autocommit = True
        with conn.cursor() as cursor:
            cursor.execute("CREATE EXTENSION IF NOT EXISTS vector")
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS embeddings (
                    id SERIAL PRIMARY KEY,
                    content TEXT,
                    metadata JSONB,
                    embedding vector(1536)
                );
            """)

def get_pg_vector_store() -> PGVector:
    """
    Return a PGVector object that references our 'embeddings' table.
    """
    embeddings = OpenAIEmbeddings(openai_api_key=st.secrets["OPENAI_API_KEY"])
    
    vector_store = PGVector(
        connection_string=CONN_STR,
        embedding_function=embeddings,
        collection_name="embeddings"
    )
    return vector_store

def extract_text_from_file(file) -> str:
    """
    Extract text from multiple file types: TXT, CSV, Markdown, PDF, DOCX, Excel.
    """
    file_type = file.name.split(".")[-1].lower()

    if file_type in ["txt", "md"]:
        return file.read().decode("utf-8", errors="ignore")

    elif file_type == "csv":
        df = pd.read_csv(file)
        return df.to_string(index=False)

    elif file_type in ["xlsx", "xls"]:
        df = pd.read_excel(file)
        return df.to_string(index=False)

    elif file_type == "docx":
        doc = docx.Document(file)
        return "\n".join([para.text for para in doc.paragraphs])

    elif file_type == "pdf":
        with pdfplumber.open(file) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

    else:
        raise ValueError(f"Unsupported file format: {file_type}")

def convert_files_to_docs(uploaded_files) -> List[Document]:
    """
    Convert uploaded files into chunked Document objects.
    """
    docs = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=50)

    for file in uploaded_files:
        try:
            content = extract_text_from_file(file)
            chunks = text_splitter.split_text(content)
            for chunk in chunks:
                docs.append(Document(page_content=chunk, metadata={"filename": file.name}))
        except Exception as e:
            st.error(f"Error processing {file.name}: {str(e)}")

    return docs

def store_embeddings_in_neon(docs: List[Document]):
    """
    Store Documents in the 'embeddings' table via PGVector.
    """
    init_rag_db()  # Ensure table exists
    vector_store = get_pg_vector_store()
    vector_store.add_documents(docs)

def create_rag_chain(
    openai_api_key: str,
    temperature: float = 0.0,
    max_tokens: int = 500,
    top_k: int = 3
) -> ConversationalRetrievalChain:
    """
    Build a Retrieval-Augmented Generation (RAG) chain that uses PGVector in NeonDB.
    """
    init_rag_db()  # Ensure embeddings table
    vector_store = get_pg_vector_store()
    retriever = vector_store.as_retriever(search_kwargs={"k": top_k})

    llm = ChatOpenAI(
        openai_api_key=openai_api_key,
        model_name="gpt-4",
        temperature=temperature,
        max_tokens=max_tokens
    )

    return ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        return_source_documents=True
    )
